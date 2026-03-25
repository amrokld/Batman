# tools.py
import smtplib
import random
import requests
import logging
import os
import smtplib
import ast
import math
import platform
import subprocess
import re
from typing import Literal
from urllib.parse import urlparse

import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="whisper")
import whisper  # must come AFTER warnings filter

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders




from config import setup_ffmpeg_path
setup_ffmpeg_path()


logger = logging.getLogger(__name__)

# Whisper model (lazy loaded)
WHISPER_MODEL = None

# --------------------calculator-----------------------------------------------------------------------------

_ALLOWED_NAMES = {
    # constants
    "pi": math.pi,
    "e": math.e,

    # math (degrees-based trig)
    "sin": lambda x: math.sin(math.radians(x)),
    "cos": lambda x: math.cos(math.radians(x)),
    "tan": lambda x: math.tan(math.radians(x)),

    # other math
    "sqrt": math.sqrt,
    "log": math.log,
    "log10": math.log10,
    "exp": math.exp,
    "abs": abs,
    "round": round,
}

_ALLOWED_NODES = (
    ast.Expression,
    ast.BinOp,
    ast.UnaryOp,
    ast.Call,
    ast.Num,
    ast.Constant,
    ast.Name,
    ast.Load,
    # operators
    ast.Add,
    ast.Sub,
    ast.Mult,
    ast.Div,
    ast.FloorDiv,
    ast.Mod,
    ast.Pow,
    ast.UAdd,
    ast.USub,
)

#--------------------Caluclate--------------------------------------------------------------
def calculate(expression: str) -> str:
    """
    Safely evaluate a math expression.
    
    """
    try:
        tree = ast.parse(expression, mode="eval")

        for node in ast.walk(tree):
            if not isinstance(node, _ALLOWED_NODES):
                raise ValueError("Unsupported expression")

            if isinstance(node, ast.Name) and node.id not in _ALLOWED_NAMES:
                raise ValueError(f"Unknown name: {node.id}")

            if isinstance(node, ast.Call):
                    if not isinstance(node.func, ast.Name) or node.func.id not in _ALLOWED_NAMES:
                      raise ValueError("Function call not allowed")

        result = eval(
            compile(tree, "<calc>", "eval"),
            {"__builtins__": {}},
            _ALLOWED_NAMES,
        )

        return str(round(result, 10)).rstrip("0").rstrip(".")

    except Exception as e:
        return f"Error: {e}"



#-----------------secret---------------------------------------------------------------------------------------------------
def get_secret_word() -> str:
    logger.info("Tool called: get_secret_word()")
    return random.choice(["Dawoooood", "Amrooooo", "fama7aga"])



#----------------weather------------------------------------------------------------------------------------------------

def get_current_weather(city: str) -> str:
    logger.info(f"Tool called: get_current_weather({city})")
    try:
        response = requests.get(f"http://wttr.in/{city}?format=3", timeout=20)
        response.raise_for_status()
        return response.text
    except requests.RequestException as e:
        logger.error(f"Weather fetch error: {e}")
        return f"Error fetching weather: {e}"
 
 
 
#----------------send email (Gmail)-----------------------------------------------------------------------------------------

def send_email_gmail(
    to: str,
    subject: str,
    body: str,
    attachment: str | None = None,
    gmail_user: str | None = None,
    gmail_password: str | None = None

) -> str:
    """
    Send an email using Gmail SMTP with optional file attachment.

    attachment:
      - filename only (e.g. "steps.pdf") → taken from written_files/
      - OR full absolute path
    """

    logger.info(f"Tool called: send_email_gmail(to={to}, subject={subject})")

    if not gmail_user or not gmail_password:
        return "Please configure your Gmail first (email + app password)."
    
    msg = MIMEMultipart()
    msg["From"] = gmail_user
    msg["To"] = to
    msg["Subject"] = subject

    msg.attach(MIMEText(body, "plain"))

    # ---------- Attachment handling ----------
    if attachment:
        # If not absolute, assume it's in written_files
        if not os.path.isabs(attachment):
            attachment = os.path.join(SAFE_WRITE_DIR, os.path.basename(attachment))

        if not os.path.exists(attachment):
            return f"Attachment not found: {attachment}"

        with open(attachment, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())

        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f'attachment; filename="{os.path.basename(attachment)}"',
        )

        msg.attach(part)

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=30) as server:
            server.login(gmail_user, gmail_password)
            server.send_message(msg)

        return f"Email sent successfully to {to}" + (
            f" with attachment {os.path.basename(attachment)}" if attachment else ""
        )

    except smtplib.SMTPAuthenticationError :
        return "Invalid Gmail credentials (check your app password)."
    except Exception as e:
        return f"Email sending failed: {e}"
    
#------------------------------------------Write files -------------------------------------------------------------  
def next_available_name(path: str) -> str:
    """
    If file exists, generate: name_1.ext, name_2.ext, ...
    """
    if not os.path.exists(path):
        return path

    base, ext = os.path.splitext(path)
    i = 1
    while True:
        new_path = f"{base}_{i}{ext}"
        if not os.path.exists(new_path):
            return new_path
        i += 1


SAFE_WRITE_DIR = os.path.join(os.path.dirname(__file__), "written_files")
os.makedirs(SAFE_WRITE_DIR, exist_ok=True)

def write_file(filename: str, content: str = "", mode: str = "write") -> str:
    """
    One tool for text/pdf/docx.
    mode: "write" (overwrite existing file) or "append".
    Note: 'append' mode is NOT supported for .pdf files.
    """
    filename = os.path.basename(filename)
    path = os.path.join(SAFE_WRITE_DIR, filename)
        
    ext = os.path.splitext(filename)[1].lower()
    mode = (mode or "write").lower().strip()

    if mode not in {"write", "append"}:
        return "Error: mode must be 'write' or 'append'."

    try:
        # ---------------- PDF ----------------
        if ext == ".pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
            except Exception:
                return "Error: PDF support missing. Install: pip install reportlab"

            # NOTE: reportlab can't truly "edit" existing PDFs easily.
            # So for append, simplest behavior: create a NEW pdf if doesn't exist,
            # otherwise rewrite a new PDF with old text + new text (we do NOT have old text).
            # Practical approach: append = add a new page visually, but we'd need a PDF merger lib.
            # For now: implement append as "add new content on a new page" by drawing AFTER a page break.
            c = canvas.Canvas(path, pagesize=A4)
            width, height = A4

            x, y = 40, height - 40
            line_h = 14

            # If append and file exists, we can't easily keep existing pages without extra libs.
            # So we warn and overwrite (or you can install pypdf to merge).
            if mode == "append" and os.path.exists(path):
                # simplest honest behavior:
                return "Error: PDF append needs a merge library. Use mode='write' or we can add PDF-merge support."

            for line in (content or "").split("\n"):
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(x, y, line)
                y -= line_h

            c.save()
            return {"ok": True, "path": path, "filename": os.path.basename(path)}

        # ---------------- DOCX ----------------
        if ext == ".docx":
            try:
                from docx import Document
            except Exception:
                return "Error: DOCX support missing. Install: pip install python-docx"

            if mode == "append" and os.path.exists(path):
                doc = Document(path)   # open existing
            else:
                doc = Document()       # new file

            for line in (content or "").split("\n"):
                doc.add_paragraph(line)

            doc.save(path)
            return {"ok": True, "path": path, "filename": os.path.basename(path)}

        # ---------------- TEXT / OTHER ----------------
        py_mode = "a" if mode == "append" else "w"
        with open(path, py_mode, encoding="utf-8") as f:
            f.write(content or "")
            if content and not content.endswith("\n"):
                f.write("\n")

        return {"ok": True, "path": path, "filename": os.path.basename(path)}

    except Exception as e:
        return f"Error writing file: {e}"
#------------------------------------------System commands ------------------------------------------------------------- 
# --- NEW: safe system commands tool (Windows) ---

_SYSTEM_ACTIONS = {
    # Network (data-only)
    "ipconfig_all": ["ipconfig", "/all"],
    "ipconfig": ["ipconfig"],
    "routes": ["route", "print"],
    "arp_table": ["arp", "-a"],
    "netstat": ["netstat", "-ano"],
    "wifi_status": ["netsh", "wlan", "show", "interfaces"],

    # Device info (data-only)
    "systeminfo": ["systeminfo"],

    # PowerShell (data-only)
    "ps_net_ipconfig": [
        "powershell",
        "-NoProfile",
        "-Command",
        "Get-NetIPConfiguration | Format-List *"
    ],
    "ps_adapters": [
        "powershell",
        "-NoProfile",
        "-Command",
        "Get-NetAdapter | Format-Table -AutoSize"
    ],
    "ps_adapter_stats": [
        "powershell",
        "-NoProfile",
        "-Command",
        "Get-NetAdapterStatistics | Format-Table -AutoSize"
    ],
}

_HOST_RE = re.compile(r"^[a-zA-Z0-9\.\-:_]+$")  # simple safe host pattern (no spaces)

def system_commands_windows(
    action: str,
    host: str | None = None,
    count: int = 4,
    max_output_chars: int = 12000,
) -> dict:
    """
    Safe Windows system commands (DATA-ONLY).
    action: one of the keys in _SYSTEM_ACTIONS, plus: "ping"
    host: required for ping
    count: ping count (1..8)
    Returns structured dict with stdout/stderr/returncode.
    """
    if platform.system().lower() != "windows":
        return {
            "ok": False,
            "error": "system_commands_windows is Windows-only.",
            "platform": platform.system()
        }

    action = (action or "").strip()

    # Special case: ping (still data-only, but does send ICMP packets)
    if action == "ping":
        if not host or not isinstance(host, str) or not _HOST_RE.match(host):
            return {"ok": False, "error": "Invalid host for ping."}
        try:
            count = int(count)
        except Exception:
            count = 4
        count = max(1, min(count, 8))
        cmd = ["ping", host, "-n", str(count)]
    else:
        if action not in _SYSTEM_ACTIONS:
            return {
                "ok": False,
                "error": "Unknown action.",
                "allowed_actions": sorted(list(_SYSTEM_ACTIONS.keys()) + ["ping"])
            }
        cmd = _SYSTEM_ACTIONS[action]

    try:
        p = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=25,
            shell=False,
        )
        out = (p.stdout or "").strip()
        err = (p.stderr or "").strip()

        # Trim output so you don't flood the model
        if len(out) > max_output_chars:
            out = out[:max_output_chars] + "\n...[truncated]"
        if len(err) > 4000:
            err = err[:4000] + "\n...[truncated]"

        return {
            "ok": True,
            "action": action,
            "command": cmd,
            "returncode": p.returncode,
            "stdout": out,
            "stderr": err,
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "action": action, "error": "Command timed out."}
    except Exception as e:
        return {"ok": False, "action": action, "error": str(e)}


def get_public_ip() -> dict:
    """
    Data-only: gets your public IP via ipify.
    """
    try:
        r = requests.get("https://api.ipify.org?format=json", timeout=12)
        r.raise_for_status()
        return {"ok": True, "public_ip": r.json().get("ip")}
    except Exception as e:
        return {"ok": False, "error": str(e)}

#----------------browse web---------------------------------------------------------------------

def browse_web(url: str, max_chars: int = 3000) -> dict:
    """
    Fetches a webpage and returns its clean text content.
    Use this when the user asks to visit, read, or summarize a URL.
    """
    logger.info(f"Tool called: browse_web(url={url})")

    # --- Validate URL ---
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return {"ok": False, "error": "Only http and https URLs are allowed."}
        if not parsed.netloc:
            return {"ok": False, "error": "Invalid URL: missing domain."}
    except Exception as e:
        return {"ok": False, "error": f"Invalid URL: {e}"}

    # --- Fetch the page ---
    try:
        from bs4 import BeautifulSoup

        headers = {"User-Agent": "Mozilla/5.0 (compatible; AI-Assistant-Browser/1.0)"}
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()

        # --- Parse and clean HTML ---
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove noise tags (scripts, styles, nav bars, etc.)
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()

        # Extract clean text
        text = soup.get_text(separator="\n", strip=True)

        # Collapse excessive blank lines
        lines = [line for line in text.splitlines() if line.strip()]
        clean_text = "\n".join(lines)

        # Truncate so we don't flood the model
        if len(clean_text) > max_chars:
            clean_text = clean_text[:max_chars] + "\n\n[Content truncated...]"

        return {
            "ok": True,
            "url": url,
            "content": clean_text,
        }

    except requests.exceptions.Timeout:
        return {"ok": False, "url": url, "error": "Request timed out."}
    except requests.exceptions.HTTPError as e:
        return {"ok": False, "url": url, "error": f"HTTP error: {e}"}
    except Exception as e:
        return {"ok": False, "url": url, "error": str(e)}

     
      
#----------------speech to text (Python whisper)---------------------------------------------------------------------

def speech_to_text(audio_path: str) -> str:
    """
    Transcribe an audio file (wav, mp3, m4a, etc.) into text using
    the local Python whisper library. No API key needed.
    """
    global WHISPER_MODEL
    logger.info(f"Tool called: speech_to_text()")

    if not os.path.exists(audio_path):
        return f"Error: audio file not found: {audio_path}"

    try:
        if WHISPER_MODEL is None:
            logger.info("Loading Whisper model (first time)...")
            WHISPER_MODEL = whisper.load_model("base")

        # Whisper automatically uses ffmpeg to read the file
        result = WHISPER_MODEL.transcribe(audio_path, language="en")
        text = result.get("text", "").strip()
        if not text:
            return "Error: transcription came back empty."
        return text
    except Exception as e:
        logger.error(f"Whisper speech_to_text error: {e}")
        return f"Error transcribing audio: {e}"

      
#---------------------------------------------------------------------------------------------------------------------------

# Registry (name → function) so the chat can look up tools by string name
TOOL_REGISTRY = {
    "calculate": calculate,
    "get_secret_word": get_secret_word,
    "get_current_weather": get_current_weather,
    "send_email_gmail": send_email_gmail,
    "speech_to_text": speech_to_text,
    "write_file": write_file,
    "system_commands_windows": system_commands_windows,
    "get_public_ip": get_public_ip,
    "browse_web": browse_web,
}

TOOL_SYSTEM_PROMPT = """
You are a helpful AI Assistant with access to external tools.

CRITICAL OUTPUT RULES (ALL MODELS):
1. If you need to use a tool, your entire response must be a SINGLE valid JSON object.
2. Do not add any text, markdown, or explanations before or after the JSON.
3. If no tool is needed, respond in plain text (no JSON).
4. ⚠️ **NO PROACTIVE ACTIONS**: ONLY use tools if the user EXPLICITLY requested that action.
   - If user says "Check my IP", you call `ipconfig` -> Then STOP and tell the user the IP. 
   - ❌ DO NOT automatically create a file or email it.
   - Only create a file if they specifically say "Save to file" or "Create document".
   - Only email if they specifically say "Email this" or "Send to...".
   - 🛡️ **EMAIL SAFETY**: If the user's intent to email is not 100% explicit and confirmed, ASK the user "Would you like me to send this email now?" before calling the tool.

--- PROACTIVE BEHAVIOR (WRONG) ---
User: "What is my network status?"
❌ BAD RESPONSE (One Turn): Call `ipconfig` AND `write_file`.
❌ BAD RESPONSE (Multi-Turn): Call `ipconfig` -> then automatically call `write_file` without asking.

--- REACTIVE BEHAVIOR (CORRECT) ---
User: "What is my network status?"
✅ GOOD: Call `system_commands_windows(action="ipconfig")`.
✅ GOOD (Next Turn): Plain text reply: "Your status is... [data]. Would you like me to save this or email it?"

TOOL CALL FORMAT:
{
  "tool": "tool_name",
  "args": {
    "arg1": "value1",
    "arg2": null  // Use null for optional missing arguments
  }
}

AVAILABLE TOOLS:
1. send_email_gmail(to: str, subject: str, body: str, attachment: str | null)
   - Sends an email via Gmail.
   - 'attachment': Optional. The exact filename from 'written_files' (e.g. "report.pdf") OR null if sending plain text.

2. write_file(filename: str, content: str, mode: str)
   - Writes text to a file.
   - 'mode': "write" (overwrite) or "append".
   - ⚠️ NOTE: 'append' is NOT supported for .pdf files. Use 'write' for PDFs.

3. calculate(expression: str)
   - Evaluates math (e.g., "sqrt(25) * 5").

4. get_current_weather(city: str)

5. get_public_ip()
   - Returns ONLY your external/public IP address (the IP the internet sees).
   - Use this ONLY when the user asks: "what's my public IP" or "my external IP".
   
6. system_commands_windows(action: str, host: str | null, count: int)
   - Windows system/network inspection commands (DATA-ONLY, safe, read-only).
   - Required argument: "action" (see list below)
   - Optional arguments: "host" (for ping), "count" (for ping, default 4)
   
   **Available actions:**
   - "ipconfig" → Basic IP config (adapter names, IPs, subnet masks)
   - "ipconfig_all" → Detailed IP config (DNS servers, MAC addresses, DHCP info)
   - "routes" → Routing table
   - "arp_table" → ARP cache
   - "netstat" → Active network connections
   - "wifi_status" → Wi-Fi adapter status
   - "systeminfo" → Full system information (OS, hardware, patches)
   - "ping" → Ping a host (REQUIRES "host" argument, e.g. host="google.com")
   - "ps_net_ipconfig" → PowerShell Get-NetIPConfiguration
   - "ps_adapters" → PowerShell Get-NetAdapter
   - "ps_adapter_stats" → PowerShell Get-NetAdapterStatistics

7. speech_to_text(audio_path: str)

8. get_secret_word()

9. browse_web(url: str)
   - Fetches and reads the text content of any public webpage.
   - Use this when the user says: "open this link", "visit", "browse", "read this URL", or "what does this page say".
   - Returns the cleaned text from the page.
   - Example: {"tool": "browse_web", "args": {"url": "https://example.com"}}


SPECIFIC USAGE GUIDES:

--- NETWORK/SYSTEM COMMANDS RULES ---
When the user asks about network or system information, use system_commands_windows:

USER QUERY → CORRECT TOOL CALL:
- "What is my IP configuration?" → {"tool": "system_commands_windows", "args": {"action": "ipconfig"}}
- "Show detailed IP config" → {"tool": "system_commands_windows", "args": {"action": "ipconfig_all"}}
- "Show network adapters" → {"tool": "system_commands_windows", "args": {"action": "ipconfig"}}
- "Ping google.com" → {"tool": "system_commands_windows", "args": {"action": "ping", "host": "google.com", "count": 4}}
- "System information" → {"tool": "system_commands_windows", "args": {"action": "systeminfo"}}
- "Show routing table" → {"tool": "system_commands_windows", "args": {"action": "routes"}}
- "What's my public IP?" → {"tool": "get_public_ip", "args": {}}

CRITICAL:
- NEVER hallucinate action names. Use ONLY the exact actions listed above.
- For ping, you MUST provide "host" argument.
- Do NOT use ipconfig flags in args (e.g., no "/all"). Use action="ipconfig_all" instead.

--- EMAIL RULES ---
Case A: Sending a File
- If the user says "email this file" or "send the pdf":
- You MUST explicitly set "attachment": "filename.ext".

Case B: Sending Plain Text (Hello, Updates, etc.)
- If the user says "email bob saying hello" or "send an email":
- You MUST explicitly set "attachment": null.
- Infer a subject if one is not given (e.g., "Hello").

--- MATH RULES ---
- ALWAYS use 'calculate' for math. Do not calculate in your head.

--- FILE WRITING RULES ---
USER QUERY → CORRECT TOOL CALL:
- "Create file.txt with Hello" → {"tool": "write_file", "args": {"filename": "file.txt", "content": "Hello", "mode": "write"}}
- "Write shopping list to list.txt" → {"tool": "write_file", "args": {"filename": "list.txt", "content": "milk\neggs\nbread", "mode": "write"}}
- "Append 'Task 3' to file.txt" → {"tool": "write_file", "args": {"filename": "file.txt", "content": "Task 3", "mode": "append"}}
- "Create report.pdf with summary" → {"tool": "write_file", "args": {"filename": "report.pdf", "content": "...", "mode": "write"}}
- After write_file succeeds, always show the clickable download link (e.g., [Download filename.pdf](http://...)) if the download_url is provided in the tool result. 
- ⚠️ NEVER use image tags ![]() for these download links. Use standard [Text](URL) format.

CRITICAL RULES:
1. For mode="write": TRUE overwrite of the existing file. (Provide FULL content).
2. For mode="append": Add ONLY the new content to the end of the file (do NOT include existing content).
3. When appending, the "content" arg should contain ONLY what to add, not the entire file.
4. Default mode is "write" if not specified.
5. Always include complete content in a single call — don't split across multiple calls.
6. ONLY these 3 parameters exist for write_file: filename, content, mode. Do NOT use "overwrite", "path", or any other parameters!

--- MULTI-STEP WORKFLOW RULES (ONLY for multi-action requests) ---
ONLY use these rules if the user asks for MULTIPLE ACTIONS (e.g., "Check data AND save it AND email it").
When a task requires using the OUTPUT of one tool as INPUT to another tool, you MUST break it into MULTIPLE CONVERSATION TURNS.

--- BROWSING RULES ---
USER QUERY → CORRECT TOOL CALL:
- "Open https://example.com" → {"tool": "browse_web", "args": {"url": "https://example.com"}}
- "What does this page say: https://..." → {"tool": "browse_web", "args": {"url": "https://..."}}
- "Summarize this article: [url]" → {"tool": "browse_web", "args": {"url": "[url]"}}

CRITICAL:
- ONLY browse URLs that the user explicitly provides. NEVER guess or invent URLs.
- After browsing, summarize the content in plain text. Do NOT output the raw content dump.


**WRONG APPROACH (All tools in one turn with empty content):**
User: "Get system info, save to file, and email it"
❌ BAD: Call all 3 tools at once:
   1. system_commands_windows(action="systeminfo")
   2. write_file(filename="systeminfo.txt", content="", mode="write")  ← EMPTY!
   3. send_email_gmail(..., attachment="systeminfo.txt")

**CORRECT APPROACH (Sequential execution across turns):**
User: "Get system info, save to file, and email it"
✅ GOOD - Turn 1: Only call system_commands_windows, then STOP and explain:
   "I've retrieved the system information. In my next response, I'll save it to a file and email it to you."
   
Then the system will call you again, and you can:
✅ GOOD - Turn 2: Use the data from Turn 1 to write file and email:
   EXACT JSON FORMAT:
   {
     "tools": [
       {"tool": "write_file", "args": {"filename": "systeminfo.txt", "content": "<actual system data>", "mode": "write"}},
       {"tool": "send_email_gmail", "args": {"to": "user@example.com", "subject": "System Info", "body": "", "attachment": "systeminfo.txt"}}
     ]
   }

**MORE EXAMPLES WITH EXACT JSON:**

Example 1: "Ping google.com and save results to ping.txt then email it"
- Turn 1: {"tool": "system_commands_windows", "args": {"action": "ping", "host": "google.com"}}
- Turn 2: {
    "tools": [
      {"tool": "write_file", "args": {"filename": "ping.txt", "content": "<ping results from turn 1>", "mode": "write"}},
      {"tool": "send_email_gmail", "args": {"to": "user@example.com", "subject": "Ping Results", "body": "", "attachment": "ping.txt"}}
    ]
  }

Example 2: "Check network config, write to network.txt, and send to user@example.com"
- Turn 1: {"tool": "system_commands_windows", "args": {"action": "ipconfig_all"}}
- Turn 2: {
    "tools": [
      {"tool": "write_file", "args": {"filename": "network.txt", "content": "<network info from turn 1>", "mode": "write"}},
      {"tool": "send_email_gmail", "args": {"to": "user@example.com", "subject": "Network Config", "body": "", "attachment": "network.txt"}}
    ]
  }

Example 3: "Create report.txt with 'Status: OK' and email it" (NO dependency - file content is given)
- Single Turn: {
    "tools": [
      {"tool": "write_file", "args": {"filename": "report.txt", "content": "Status: OK", "mode": "write"}},
      {"tool": "send_email_gmail", "args": {"to": "user@example.com", "subject": "Report", "body": "", "attachment": "report.txt"}}
    ]
  }
  ↑ This is OK because content is provided by user, not from another tool


**KEY PRINCIPLE:**
- If you need tool OUTPUT → Use it as INPUT → Execute in SEPARATE turns
- If all inputs are known upfront → Execute in SINGLE turn

Reminder:
- For 7B Models: Do not hallucinate arguments. If 'attachment' is not needed, send null.
- For 120B Models: Be precise and concise. Break complex multi-step tasks into sequential turns.
"""



# ------------------------- Helper to record audio ------------------------------------------------------
def record_from_mic(duration: int = 7, samplerate: int = 16000) -> str:
    """
     Record audio from the microphone for `duration` seconds,
    save it into ./recordings/, and return the file path.
    """
    import sounddevice as sd
    from scipy.io.wavfile import write
    import time

    # Ensure recordings directory exists
    base_dir = os.path.dirname(os.path.abspath(__file__))
    recordings_dir = os.path.join(base_dir, "recordings")
    os.makedirs(recordings_dir, exist_ok=True)

    filename = f"talk_{int(time.time())}.wav"
    file_path = os.path.join(recordings_dir, filename)

    print(f"🎤 Listening... ({duration} seconds)")
    audio = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=1)
    sd.wait()  # wait until recording is finished
    write(file_path, samplerate, audio)  # save as WAV
    print(f"✔ Saved recording.\n")

    return file_path


# -------------------Helper for read files ------------------------------------------------------------
def pick_file_from_window(initialdir=None) -> str | None:
    from tkinter import Tk, filedialog
    try:
        root = Tk()
    except Exception:
        return None
        
    root.withdraw()
    root.attributes("-topmost", True)  # bring dialog to front
    root.update()
    file_path = filedialog.askopenfilename(
        initialdir=initialdir or os.path.expanduser("~"),
        title="Select a file to send to AI"
    )
    root.destroy()
    return file_path or None


def read_files(path: str, max_chars: int = 12000, max_pages: int = 3) -> str:
    from pypdf import PdfReader
    
    ext = os.path.splitext(path)[1].lower()

    # ---- PDF ----
    if ext == ".pdf":
        try:
            reader = PdfReader(path)
            pages = min(len(reader.pages), max_pages)
            text_parts = []
            for i in range(pages):
                page_text = reader.pages[i].extract_text() or ""
                text_parts.append(f"\n--- Page {i+1} ---\n{page_text}")
            text = "\n".join(text_parts).strip()

            if not text:
                return "Error: Could not extract text from this PDF (it may be scanned images)."

            return text[:max_chars] + ("\n\n[Truncated]" if len(text) > max_chars else "")
        except Exception as e:
            return f"Error reading PDF: {e}"

    # ---- TEXT (fallback) ----
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            data = f.read(max_chars + 1)
        if len(data) > max_chars:
            data = data[:max_chars] + "\n\n[Truncated]"
        return data
    except Exception as e:
        return f"Error reading file: {e}"